# --- Final Production Code ---

import io
import os
import requests
import concurrent.futures
import logging
import time
from typing import List, Dict, Any, Optional
from functools import lru_cache

# Qdrant
import qdrant_client
from qdrant_client.http.models import Distance, VectorParams, PointStruct, CollectionStatus, OptimizersConfigDiff, HnswConfigDiff, SparseVectorParams

# LangChain components
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import Qdrant, FastEmbedSparse, RetrievalMode, QdrantVectorStore
from langchain_openai import OpenAIEmbeddings # Use OpenAI embeddings

# File Parsers
import pdfplumber
from docx import Document as DocxDocument

# Add these imports at the top of the file
from fastapi import WebSocket
from fastapi.responses import StreamingResponse
import asyncio
import aiohttp
import json
import sys

# Setup robust logging for production
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("kb_indexer.log"), # Log to a file
        logging.StreamHandler() # Log to console
    ]
)
logger = logging.getLogger(__name__)

# --- Constants ---
# Dimension for OpenAI's text-embedding-3-small
OPENAI_EMBEDDING_DIMENSION = 1536
# Default Qdrant distance metric suitable for OpenAI embeddings
QDRANT_DISTANCE_METRIC = Distance.COSINE
# Default sparse embedding model
DEFAULT_SPARSE_MODEL = "Qdrant/bm25"

# --- Helper Function to Parse Files ---
def parse_file_content(content: bytes, filename: str) -> Optional[str]:
    """Parses content based on filename extension."""
    logger.debug(f"Parsing file: {filename}")
    try:
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext == '.pdf':
            text = ""
            # Using BytesIO for in-memory processing
            with io.BytesIO(content) as pdf_stream:
                with pdfplumber.open(pdf_stream) as pdf:
                    # Handle potential empty pages or extract errors gracefully
                    for page_num, page in enumerate(pdf.pages):
                         page_text = page.extract_text()
                         if page_text:
                             text += page_text + "\n"
                         else:
                             logger.debug(f"No text extracted from page {page_num + 1} of {filename}")
            logger.debug(f"Finished parsing PDF: {filename}")
            return text.strip() if text else None # Return None if empty
        elif file_ext == '.docx':
             # Using BytesIO for in-memory processing
            with io.BytesIO(content) as docx_stream:
                doc = DocxDocument(docx_stream)
                text = "\n".join(para.text for para in doc.paragraphs if para.text)
            logger.debug(f"Finished parsing DOCX: {filename}")
            return text.strip() if text else None # Return None if empty
        elif file_ext == '.txt':
            # Attempt decoding with UTF-8, replace errors
            text = content.decode('utf-8', errors='replace')
            logger.debug(f"Finished parsing TXT: {filename}")
            return text.strip() if text else None # Return None if empty
        else:
            logger.warning(f"Unsupported file type: {filename}. Skipping.")
            return None
    except Exception as e:
        logger.error(f"Error parsing file {filename}: {e}", exc_info=True) # Log traceback
        return None

# --- Helper Function to Process a Single File ---
def process_single_file(file_url: str, splitter: RecursiveCharacterTextSplitter) -> List[Document]:
    """Downloads, parses, and chunks a single file."""
    # Check if it's a valid URL (only support http/https URLs)
    if not file_url.startswith(('http://', 'https://')):
        logger.warning(f"Invalid URL format: {file_url}")
        return []
        
    # Rest of the existing function for remote files
    documents = []
    try:
        logger.info(f"Processing remote file: {file_url}")
        # Increased timeout for potentially large files - INCREASE TIMEOUT
        response = requests.get(file_url, timeout=30)  # Reduced from 120 to 30 seconds
        response.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)

        # More robust filename extraction from URL
        try:
            # Attempt to get filename from Content-Disposition header first
            content_disposition = response.headers.get('Content-Disposition')
            if content_disposition:
                import re
                fname = re.findall('filename="?(.+)"?', content_disposition)
                if fname:
                    filename = fname[0]
                else: # Fallback if header is malformed
                    filename = os.path.basename(file_url.split('?')[0].split('/')[-1])

            else: # Fallback to URL path component
                 filename = os.path.basename(file_url.split('?')[0].split('/')[-1])

            # Handle cases where filename might be empty or just '/'
            if not filename or filename == '/':
                filename = f"unknown_file_{time.time()}" # Generate a placeholder
                logger.warning(f"Could not determine filename for {file_url}, using placeholder: {filename}")

        except Exception as e:
             logger.warning(f"Error extracting filename for {file_url}, using fallback. Error: {e}")
             filename = f"unknown_file_{time.time()}" # Generate a placeholder


        content = response.content
        if not content:
             logger.warning(f"Downloaded file {filename} ({file_url}) is empty. Skipping.")
             return []

        extracted_text = parse_file_content(content, filename)

        if extracted_text:
            # Use the splitter to create chunks
            chunks = splitter.split_text(extracted_text)
            # Create LangChain Document objects with metadata
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_url,
                        "filename": filename,
                        "chunk_index": i,
                        "indexed_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()) # Add timestamp
                    }
                )
                documents.append(doc)
            logger.info(f"Successfully chunked {filename} into {len(documents)} documents.")
        else:
             logger.warning(f"No text could be extracted from {filename} ({file_url}).")


    except requests.exceptions.Timeout:
        logger.error(f"Timeout while downloading {file_url}")
    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to download or access {file_url}: {e}")
    except Exception as e:
        # Catch any other unexpected errors during processing
        logger.error(f"An unexpected error occurred processing {file_url}: {e}", exc_info=True)

    return documents

# --- Main Indexer Function ---
def KB_indexer(
    file_urls: List[str],
    qdrant_url: str,
    qdrant_api_key: Optional[str],
    collection_name: str,
    openai_api_key: Optional[str] = None,
    force_recreate_collection: bool = False,
    max_workers: int = 15,
    use_hybrid_search: bool = False,
    system_prompt: Optional[str] = None,
    schema: Optional[Dict] = None,
    fast_mode: bool = False,
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> bool:
    start_time = time.time()
    logger.info(f"Starting indexing process for collection: {collection_name}, fast_mode: {fast_mode}")
    
    # OPTIMIZATION 1: Use OpenAI's ada embeddings for speed in fast mode
    embedding_model_name = "text-embedding-3-small"
    if fast_mode:
        # Much faster than the default model
        embedding_model_name = "text-embedding-ada-002"
    
    # Initialize the embedding model with caching
    embedding_model = OpenAIEmbeddings(
        api_key=openai_api_key or os.environ.get("OPENAI_API_KEY"),
        model=embedding_model_name
    )
    
    # OPTIMIZATION 2: Skip sparse embeddings in fast mode
    if fast_mode and use_hybrid_search:
        logger.info("Fast mode enabled - using dense vectors only for quick processing")
        use_hybrid_search = False  # Disable hybrid search for speed
    
    # OPTIMIZATION 3: Adjusted text splitter parameters
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        # Use faster separators
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # OPTIMIZATION 4: Smart parallel processing
    # Process files in parallel with adaptive chunking
    all_documents = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Process files in batches for better memory management
        batch_size = 5  # Process 5 files at a time
        
        for i in range(0, len(file_urls), batch_size):
            batch_urls = file_urls[i:i+batch_size]
            future_to_url = {executor.submit(process_single_file, url, splitter): url for url in batch_urls}
            
            for future in concurrent.futures.as_completed(future_to_url):
                url = future_to_url[future]
                try:
                    docs = future.result()
                    all_documents.extend(docs)
                except Exception as e:
                    logger.error(f"Error processing {url}: {e}")

    logger.info(f"Processed {len(all_documents)} document chunks from {len(file_urls)} files")
    
    # OPTIMIZATION 5: Batch embedding and upsert
    if all_documents:
        try:
            # Initialize Qdrant client with optimized settings
            client = qdrant_client.QdrantClient(
                url=qdrant_url,
                api_key=qdrant_api_key,
                timeout=30.0  # Reduced timeout for faster operation
            )
            
            # Check if collection exists
            try:
                collection_info = client.get_collection(collection_name)
                collection_exists = True
            except Exception:
                collection_exists = False
            
            # Create or recreate collection
            if force_recreate_collection and collection_exists:
                client.delete_collection(collection_name)
                collection_exists = False

            if not collection_exists:
                if use_hybrid_search:
                    # Configure both dense and sparse vectors for hybrid search
                    vectors_config = {
                        "dense": VectorParams(
                            size=OPENAI_EMBEDDING_DIMENSION,
                            distance=QDRANT_DISTANCE_METRIC,
                            name="dense"
                        ),
                        "sparse": SparseVectorParams(
                            name="sparse"
                        )
                    }
                    client.create_collection(
                        collection_name=collection_name,
                        vectors_config=vectors_config
                    )
                else:
                    # Standard dense-only configuration
                    vectors_config = VectorParams(
                        size=OPENAI_EMBEDDING_DIMENSION,
                        distance=QDRANT_DISTANCE_METRIC
                    )
                    # Create collection with standard config
                    client.create_collection(
                        collection_name=collection_name,
                        vectors_config=vectors_config
                    )
            
            # OPTIMIZATION 6: Batch processing for embeddings
            # The optimal batch size for OpenAI API is around 100 items
            BATCH_SIZE = 100
            
            # Process in batches to avoid memory issues
            for i in range(0, len(all_documents), BATCH_SIZE):
                batch = all_documents[i:i+BATCH_SIZE]
                
                # Use vectorstore with hybrid search settings for batch operations
                if use_hybrid_search:
                    try:
                        sparse_embedding_model = FastEmbedSparse(model_name=DEFAULT_SPARSE_MODEL)
                        vectorstore = QdrantVectorStore(
                            client=client,
                            collection_name=collection_name,
                            embedding=embedding_model,
                            sparse_embedding=sparse_embedding_model,
                            retrieval_mode=RetrievalMode.HYBRID,
                            vector_name="dense",
                            sparse_vector_name="sparse"
                        )
                    except Exception as e:
                        logger.warning(f"Failed to use hybrid search during indexing: {e}. Falling back to dense embeddings.")
                        vectorstore = QdrantVectorStore(
                            client=client,
                            collection_name=collection_name,
                            embedding=embedding_model
                        )
                else:
                    vectorstore = QdrantVectorStore(
                        client=client,
                        collection_name=collection_name,
                        embedding=embedding_model
                    )

                # Add documents in batch (faster than individual operations)
                vectorstore.add_documents(batch)

                logger.info(f"Indexed batch of {len(batch)} documents ({i+len(batch)}/{len(all_documents)})")
            
            total_time = time.time() - start_time
            logger.info(f"Indexing process completed successfully in {total_time:.2f} seconds.")
            return True

        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"Indexing error: {e}", exc_info=True)
            logger.error(f"Indexing process failed after {total_time:.2f} seconds.")
            return False
    else:
        logger.warning("No documents were generated from the provided URLs.")
        return False

# --- RAG Query Functions ---
def retrieve_documents(
    query: str,
    qdrant_url: str,
    qdrant_api_key: Optional[str],
    collection_name: str,
    openai_api_key: Optional[str] = None,
    top_k: int = 5,
    prefix: str = "",
    use_hybrid_search: bool = False
) -> List[Document]:
    """
    Retrieves relevant documents from a Qdrant collection using the query.
    """
    logger.info(f"Retrieving documents from collection: {collection_name}")
    
    try:
        # Initialize client
        client = qdrant_client.QdrantClient(
            url=qdrant_url,
            api_key=qdrant_api_key,
            timeout=60.0
        )
        
        # Check if collection exists before attempting to retrieve
        try:
            client.get_collection(collection_name=collection_name)
        except qdrant_client.http.exceptions.UnexpectedResponse as e:
            # Check if it's a 404 error (collection not found)
            if hasattr(e, 'status_code') and e.status_code == 404:
                logger.warning(f"Collection {collection_name} does not exist. Returning empty results.")
                return []
            else:
                # If it's another type of error, re-raise
                raise
        
        # Continue with normal processing if collection exists
        # Initialize embeddings model
        embedding_model = OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=openai_api_key,
        )
        
        # Create vector store
        if use_hybrid_search:
            try:
                # Check if collection exists and get info about vectors
                collection_info = client.get_collection(collection_name)
                
                # More robust check for vector params - VectorParams object doesn't support len()
                has_named_vectors = False
                try:
                    # For new collections with named vectors dict structure
                    if hasattr(collection_info.config.params, 'vectors') and collection_info.config.params.vectors:
                        # If it's a dictionary of named vectors
                        if isinstance(collection_info.config.params.vectors, dict):
                            has_named_vectors = True
                    # For collections with vector_config using the 'name' property
                    elif (hasattr(collection_info.config.params, 'vectors_config') and 
                          collection_info.config.params.vectors_config and
                          hasattr(collection_info.config.params.vectors_config, 'name') and
                          collection_info.config.params.vectors_config.name):
                        has_named_vectors = True
                except Exception as e:
                    logger.debug(f"Error checking vector names: {e}")
                    
                # Configure vector names based on collection structure
                vector_name = "dense" if has_named_vectors else ""
                sparse_vector_name = "sparse" if has_named_vectors else ""
                
                logger.info(f"Initializing hybrid search with vector_name={vector_name}, sparse_vector_name={sparse_vector_name}")
                
                # Initialize hybrid search
                sparse_embedding_model = FastEmbedSparse(model_name=DEFAULT_SPARSE_MODEL)
                vector_store = QdrantVectorStore(
                    client=client,
                    collection_name=collection_name,
                    embedding=embedding_model,
                    sparse_embedding=sparse_embedding_model,
                    retrieval_mode=RetrievalMode.HYBRID,
                    vector_name=vector_name,
                    sparse_vector_name=sparse_vector_name
                )
            except Exception as e:
                logger.warning(f"Failed to initialize hybrid search: {e}. Falling back to dense search.")
                vector_store = QdrantVectorStore(
                    client=client,
                    collection_name=collection_name,
                    embedding=embedding_model
                )
        else:
            # Standard dense-only initialization
            vector_store = QdrantVectorStore(
                client=client,
                collection_name=collection_name,
                embedding=embedding_model
            )
        
        # Cache the query embedding for performance
        query_embedding = get_cached_embedding(query, openai_api_key)
        
        # Search for similar documents
        retrieved_docs = vector_store.similarity_search(query, k=top_k)
        
        # Add prefix to document content if provided
        if prefix:
            for doc in retrieved_docs:
                doc.page_content = f"{prefix} {doc.page_content}"
        
        logger.info(f"Retrieved {len(retrieved_docs)} documents from {collection_name}")
        return retrieved_docs
        
    except Exception as e:
        logger.error(f"Error retrieving documents from {collection_name}: {e}", exc_info=True)
        return []

# Add embedding caching
@lru_cache(maxsize=100)
def get_cached_embedding(text, api_key):
    embedding_model = OpenAIEmbeddings(
        api_key=api_key,
        model="text-embedding-3-small"
    )
    return embedding_model.embed_query(text)

# Modify the perform_rag_query_stream function to support multiple models
async def perform_rag_query_stream(
    query: str,
    base_collection_name: str,
    user_collection_name: Optional[str], 
    qdrant_url: str,
    qdrant_api_key: Optional[str],
    openai_api_key: Optional[str],
    model_api_key: Optional[str] = None,
    history: List[Dict[str, str]] = None,
    memory: List[Dict[str, str]] = None,
    model: str = "gpt-4o-mini",
    top_k: int = 5,  # Increased from 3 to 5
    use_hybrid_search: bool = False,
    system_prompt: Optional[str] = None  # Add system prompt parameter
) -> StreamingResponse:
    logger.info(f"Performing streaming RAG query using model: {model}")
    history = history or []
    
    # Process memory to enhance context
    memory_context = ""
    if memory and len(memory) > 0:
        # Format recent memory items (limited to last 10 to avoid token bloat)
        recent_memory = memory[-10:] if len(memory) > 10 else memory
        memory_context = "Previous conversation context:\n"
        for i, msg in enumerate(recent_memory):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            memory_context += f"[{role}]: {content}\n"
        memory_context += "\n"
    
    # Determine provider based on model name
    provider = "openai"  # Default provider
    if model.startswith("claude"):
        provider = "anthropic"
    elif model.startswith("gemini"):
        provider = "google"
    elif model.startswith("llama"):
        provider = "meta"
        
    logger.info(f"Using {provider} provider for model: {model}")
    
    try:
        # Get documents from both collections if they exist
        base_docs = []
        try:
            base_docs = retrieve_documents(
                query=query,
                qdrant_url=qdrant_url,
                qdrant_api_key=qdrant_api_key,
                collection_name=base_collection_name,
                openai_api_key=openai_api_key,
                top_k=top_k,
                prefix="Knowledge Base",
                use_hybrid_search=use_hybrid_search
            )
        except Exception as e:
            logger.warning(f"Error retrieving from base collection: {str(e)}")
        
        user_docs = []
        if user_collection_name:
            try:
                user_docs = retrieve_documents(
                    query=query,
                    qdrant_url=qdrant_url,
                    qdrant_api_key=qdrant_api_key,
                    collection_name=user_collection_name,
                    openai_api_key=openai_api_key,
                    top_k=top_k,
                    prefix="User Uploaded",
                    use_hybrid_search=use_hybrid_search
                )
            except Exception as e:
                logger.warning(f"Error retrieving from user collection: {str(e)}")
        
        # Combine docs, prioritizing user docs
        all_docs = user_docs + base_docs
        
        # Prepare context from retrieved documents without additional instructions
        context = ""
        if all_docs:
            context = "RELEVANT INFORMATION:\n\n"
            for i, doc in enumerate(all_docs[:top_k*2]):
                source = doc.metadata.get("source", "Unknown")
                filename = doc.metadata.get("filename", "")
                
                # Format document content for better readability
                content = doc.page_content.strip()
                
                # Add clear document source markers that will be preserved in the final output
                context += f"=== DOCUMENT {i+1}: {filename or source} ===\n{content}\n\n"
        
        # Clean up the message sequence - this is what's causing the problem
        messages = []
        
        # Use the system prompt from frontend directly
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        # No else statement - only use system prompt if provided
        
        # Add context to the user's actual query instead of as a separate message
        if context:
            user_query_with_context = f"""
{query}

Here is the ONLY relevant information to help answer my question:
{context}

IMPORTANT: Only use the information provided above to answer the question. If the information provided is insufficient to answer the question completely, explicitly state what specific information is missing. DO NOT make up or hallucinate any information like names, statistics, or details that are not explicitly stated in the provided context.

Please format your response using proper markdown.
"""
            
            # Add conversation history
            for message in history:
                messages.append({
                    "role": message["role"],
                    "content": message["content"]
                })
            
            # Add the current query with context
            messages.append({"role": "user", "content": user_query_with_context})
        else:
            # When no context is found, explicitly mention that no relevant information was found
            user_query_with_context = f"""
{query}

IMPORTANT: I could not find any relevant information in the provided documents to answer your question. 
Please let me know that you don't have the requested information rather than making up an answer.
"""
            
            # Add conversation history
            for message in history:
                messages.append({
                    "role": message["role"],
                    "content": message["content"]
                })
            
            # Add the current query with context
            messages.append({"role": "user", "content": user_query_with_context})
        
        # Streaming response
        async def streaming_response():
            async with aiohttp.ClientSession() as session:
                # Configure API endpoint and payload based on provider
                if provider == "openai":
                    api_url = "https://api.openai.com/v1/chat/completions"
                    payload = {
                        "model": model,
                        "messages": messages,
                        "stream": True,
                        "temperature": 0.4,  # Slightly higher temperature for more creative markdown formatting
                        "top_p": 0.92,       # Allow some creativity while staying focused
                        "presence_penalty": 0.1,  # Slight preference for new content
                        "frequency_penalty": 0.2,  # Reduce repetition slightly
                        "response_format": {"type": "text"}  # Ensure we get proper text, not JSON
                    }
                    headers = {
                        "Authorization": f"Bearer {openai_api_key}",
                        "Content-Type": "application/json"
                    }
                else:
                    # Default to OpenAI API for all other providers for now
                    api_url = "https://api.openai.com/v1/chat/completions"
                    fallback_model = "gpt-4o-mini"  # Use a reliable fallback model
                    logger.info(f"Using OpenAI fallback model {fallback_model} instead of {provider}")
                    
                    payload = {
                        "model": fallback_model,
                        "messages": messages,
                        "stream": True,
                        "temperature": 0.4,  # Slightly higher temperature for more creative markdown formatting
                        "top_p": 0.92,       # Allow some creativity while staying focused
                        "presence_penalty": 0.1,  # Slight preference for new content
                        "frequency_penalty": 0.2,  # Reduce repetition slightly
                        "response_format": {"type": "text"}  # Ensure we get proper text, not JSON
                    }
                    headers = {
                        "Authorization": f"Bearer {openai_api_key}",
                        "Content-Type": "application/json"
                    }
                    
                async with session.post(api_url, headers=headers, json=payload) as response:
                    if not response.ok:
                        error_msg = f"API error: {response.status} - {await response.text()}"
                        logger.error(error_msg)
                        yield f"data: {json.dumps({'error': error_msg})}\n\n"
                        yield f"data: {json.dumps({'done': True})}\n\n"
                        return
                        
                    async for line in response.content:
                        if line.startswith(b"data: "):
                            line = line[6:].strip()
                            if line == b"[DONE]":
                                yield f"data: {json.dumps({'done': True})}\n\n"
                                break
                            try:
                                data = json.loads(line)
                                if "choices" in data and len(data["choices"]) > 0:
                                    delta = data["choices"][0].get("delta", {})
                                    content = delta.get("content", "")
                                    if content:
                                        yield f"data: {json.dumps({'content': content})}\n\n"
                            except json.JSONDecodeError:
                                logger.error(f"Error decoding JSON: {line}")
                                # Skip this line instead of failing
                                continue
        return StreamingResponse(streaming_response(), media_type="text/event-stream")
    except Exception as e:
        logger.error(f"Error in perform_rag_query_stream: {e}", exc_info=True)
        return StreamingResponse(content=f"data: {json.dumps({'error': str(e)})}\n\ndata: {json.dumps({'done': True})}\n\n", media_type="text/event-stream")

# Similarly update the non-streaming function
def perform_rag_query(
    query: str,
    base_collection_name: str,
    user_collection_name: Optional[str], 
    qdrant_url: str,
    qdrant_api_key: Optional[str],
    openai_api_key: Optional[str],
    history: List[Dict[str, str]] = None,
    memory: List[Dict[str, str]] = None,
    model: str = "gpt-4o-mini",
    use_hybrid_search: bool = False,
    system_prompt: Optional[str] = None  # Add system prompt parameter
) -> str:
    # Process memory to enhance context
    memory_context = ""
    if memory and len(memory) > 0:
        # Format recent memory items (limited to last 10 to avoid token bloat)
        recent_memory = memory[-10:] if len(memory) > 10 else memory
        memory_context = "Previous conversation context:\n"
        for i, msg in enumerate(recent_memory):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            memory_context += f"[{role}]: {content}\n"
        memory_context += "\n"
    
    try:
        # Get documents from both collections if they exist
        base_docs = []
        try:
            base_docs = retrieve_documents(
                query=query,
                qdrant_url=qdrant_url,
                qdrant_api_key=qdrant_api_key,
                collection_name=base_collection_name,
                openai_api_key=openai_api_key,
                top_k=3,
                prefix="Knowledge Base",
                use_hybrid_search=use_hybrid_search
            )
        except Exception as e:
            logger.warning(f"Error retrieving from base collection: {str(e)}")
        
        user_docs = []
        if user_collection_name:
            try:
                user_docs = retrieve_documents(
                    query=query,
                    qdrant_url=qdrant_url,
                    qdrant_api_key=qdrant_api_key,
                    collection_name=user_collection_name,
                    openai_api_key=openai_api_key,
                    top_k=3,
                    prefix="User Uploaded",
                    use_hybrid_search=use_hybrid_search
                )
            except Exception as e:
                logger.warning(f"Error retrieving from user collection: {str(e)}")
        
        # Combine docs, prioritizing user docs
        all_docs = user_docs + base_docs
        
        # Prepare context from retrieved documents with better formatting
        context = ""
        if all_docs:
            context = "RELEVANT INFORMATION:\n\n"
            for i, doc in enumerate(all_docs[:top_k*2]):
                source = doc.metadata.get("source", "Unknown")
                filename = doc.metadata.get("filename", "")
                
                # Format document content for better readability
                content = doc.page_content.strip()
                
                # Add clear document source markers that will be preserved in the final output
                context += f"=== DOCUMENT {i+1}: {filename or source} ===\n{content}\n\n"
        
        # Format the conversation history with only the original system prompt
        system_message = system_prompt if system_prompt else ""
            
        # Add context information as part of the user query, not in the system message
        context_query = f"""
Here is ALL the relevant information to help answer my question:

{context}

My question is: {query}

IMPORTANT: Only use the information provided above to answer the question. If the information provided is insufficient to answer the question completely, explicitly state what specific information is missing. DO NOT make up or hallucinate any information like names, statistics, or details that are not explicitly stated in the provided context.

Please format your response using proper markdown with:
- Headings (##, ###) for clear section organization
- Bullet points and numbered lists where appropriate
- Tables for structured comparisons if needed
- Bold and italic text for emphasis
"""
        
        # Format the conversation history
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})

        # Add history
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})

        # Add current message with context
        messages.append({"role": "user", "content": context_query})
        
        # Call language model with faster parameters
        from openai import OpenAI
        client = OpenAI(api_key=openai_api_key)
        
        # Generate response with reduced tokens for speed
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.4,  # Slightly higher temperature for more creative markdown
            max_tokens=1000,  # Increase max tokens to allow for richer formatting
            top_p=0.92,       # Allow creativity while staying focused
            presence_penalty=0.1,
            frequency_penalty=0.2,
            response_format={"type": "text"}  # Ensure we get proper text, not JSON
        )
        
        answer = response.choices[0].message.content
        logger.info("RAG query completed successfully")
        return answer
        
    except Exception as e:
        logger.error(f"Error in RAG query: {e}", exc_info=True)
        return f"I encountered an error processing your question. Please try again with a more specific query."

# 3. Add a parallel retrieval process for faster document fetching
async def retrieve_documents_parallel(query, base_collection, user_collection=None, top_k=3):
    """Retrieve documents from multiple collections in parallel"""
    async def get_collection_docs(collection_name, prefix=""):
        # Existing document retrieval code
        pass
        
    # Run both retrievals in parallel
    tasks = [
        get_collection_docs(base_collection, "Knowledge Base"),
    ]
    
    if user_collection:
        tasks.append(get_collection_docs(user_collection, "User Document"))
        
    results = await asyncio.gather(*tasks)
    
    # Combine and deduplicate results
    all_docs = []
    seen_texts = set()
    
    # Prioritize user documents
    for docs in results:
        for doc in docs:
            if doc.page_content not in seen_texts:
                seen_texts.add(doc.page_content)
                all_docs.append(doc)
                
    return all_docs[:top_k]

# --- How to Use (Instructions incorporated in comments and example) ---

if __name__ == '__main__':

    print("--- KB Indexer Example ---")

    # --- 1. Prerequisites ---
    #    - Python 3.8+ installed
    #    - Access to a Qdrant instance (local or cloud)
    #    - An OpenAI API Key

    # --- 2. Installation ---
    #    Run in your terminal:
    #    pip install qdrant-client "langchain>=0.1.0" langchain-community langchain-text-splitters langchain-openai openai tiktoken pdfplumber python-docx requests fastembed

    # --- 3. Environment Variables ---
    #    Set your OpenAI API Key. It's recommended to use environment variables for secrets.
    #    Linux/macOS: export OPENAI_API_KEY='your-openai-api-key'
    #    Windows (cmd): set OPENAI_API_KEY=your-openai-api-key
    #    Windows (PowerShell): $env:OPENAI_API_KEY='your-openai-api-key'
    #    Alternatively, you can pass the key directly to the KB_indexer function via the `openai_api_key` argument,
    #    but using environment variables is generally more secure.

    # --- 4. Configuration ---
    # Replace these placeholders with your actual values

    # Qdrant Configuration
    QDRANT_URL = os.environ.get("QDRANT_URL", "http://localhost:6333") # Your Qdrant instance URL
    QDRANT_API_KEY = os.environ.get("QDRANT_API_KEY", None) # Your Qdrant API Key (if using Qdrant Cloud or authentication)

    # File URLs to index (replace with URLs accessible by the script)
    # Ensure these are direct links to PDF, DOCX, or TXT files.
    FILE_URLS = [
        "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf", # Example PDF
        "https://www.unm.edu/~unmvclib/powerpoint/pptexamples/sample.docx", # Example DOCX (check if link works)
        "https://raw.githubusercontent.com/google/gemini-api/main/README.md", # Example TXT/MD (parsed as TXT)
        "https://invalid-url-that-will-fail.xyz/document.pdf" # Example of a failing URL
    ]

    # Collection Name Construction (example based on prompt)
    USER_EMAIL = "test@example.com"
    GPT_NAME = "my_openai_rag_gpt"
    # Sanitize email and GPT name to create a valid collection name
    # Qdrant collection names must start with a letter and contain only letters, numbers, and underscores.
    sanitized_email = ''.join(c if c.isalnum() else '_' for c in USER_EMAIL)
    sanitized_gpt_name = ''.join(c if c.isalnum() else '_' for c in GPT_NAME)
    # Ensure it starts with a letter if the sanitized email doesn't
    collection_prefix = "kb" if not sanitized_email or not sanitized_email[0].isalpha() else ""
    dynamic_collection_name = f"{collection_prefix}_{sanitized_email}_{sanitized_gpt_name}"
    # Ensure name is not excessively long if needed
    dynamic_collection_name = dynamic_collection_name[:63] # Example length limit if necessary

    # --- 5. Running the Indexer ---
    print(f"Target Qdrant URL: {QDRANT_URL}")
    print(f"Target Collection Name: {dynamic_collection_name}")
    print(f"Number of files to process: {len(FILE_URLS)}")

    # Check if OpenAI API key is available (optional check, the function handles it)
    openai_key_present = bool(os.environ.get("OPENAI_API_KEY"))
    if not openai_key_present:
         print("\nWARNING: OPENAI_API_KEY environment variable not found. Make sure it's set.")
         # You could add logic here to prompt for the key or read from a config file if needed
    else:
        print("OpenAI API key found in environment.")

    # Execute the main function
    # Set force_recreate_collection=True ONLY if you want to clear the collection first.
    # Set force_recreate_collection=False to append/update documents.
    # Set use_hybrid_search=True to enable hybrid search with sparse embeddings
    success = KB_indexer(
        file_urls=FILE_URLS,
        qdrant_url=QDRANT_URL,
        qdrant_api_key=QDRANT_API_KEY,
        collection_name=dynamic_collection_name,
        openai_api_key=None, # Set to your key string if not using env var, e.g., "sk-..."
        force_recreate_collection=True, # Be careful with True in production!
        max_workers=5, # Adjust based on your machine/network
        use_hybrid_search=True, # Enable hybrid search capability
        system_prompt="This is a system prompt",
        schema={"model": "gpt-4o-mini"},
        fast_mode=True,  # Enable fast mode
        chunk_size=1000,  # Set custom chunk size
        chunk_overlap=50  # Set custom chunk overlap
    )

    # --- 6. Result ---
    if success:
        print(f"\n--- KB Indexer finished successfully for collection '{dynamic_collection_name}' ---")
        
        # Example query using hybrid search
        print("\n--- Testing hybrid search query ---")
        try:
            result = perform_rag_query(
                query="What does this document talk about?",
                base_collection_name=dynamic_collection_name,
                user_collection_name=None,
                qdrant_url=QDRANT_URL,
                qdrant_api_key=QDRANT_API_KEY,
                openai_api_key=None,
                use_hybrid_search=True
            )
            print(f"Query result: {result}")
        except Exception as e:
            print(f"Error testing query: {e}")
    else:
        print(f"\n--- KB Indexer failed for collection '{dynamic_collection_name}'. Check 'kb_indexer.log' for details. ---")